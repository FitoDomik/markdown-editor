import sys
import os
import re
import json
import markdown
import datetime
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QSplitter, QPlainTextEdit, QTextEdit,
    QTextBrowser, QVBoxLayout, QWidget, QToolBar, QMenu,
    QMenuBar, QStatusBar, QFileDialog, QMessageBox, QDialog,
    QGridLayout, QLabel, QLineEdit, QPushButton, QCheckBox, QGroupBox,
    QTreeView, QDockWidget, QInputDialog, QHBoxLayout
)
from PyQt6.QtCore import (
    Qt, QSize, QTimer, pyqtSignal as Signal, pyqtSlot as Slot, QSettings, QRect, QRegularExpression,
    QDir, QFileInfo, QModelIndex, QItemSelectionModel, QFile, QTextStream
)
from PyQt6.QtGui import (
    QAction, QIcon, QFont, QColor, QTextCharFormat, QPainter, 
    QTextCursor, QSyntaxHighlighter, QTextFormat, QPalette, QTextDocument,
    QFileSystemModel, QStandardItemModel, QStandardItem, QKeySequence, QPixmap, QGuiApplication, QImage, QShortcut
)
try:
    from markdown_it import MarkdownIt
    from pygments import highlight
    from pygments.lexers import get_lexer_by_name
    from pygments.formatters import HtmlFormatter
    MARKDOWN_IT_AVAILABLE = True
except ImportError:
    MARKDOWN_IT_AVAILABLE = False
    print("Warning: markdown-it-py or pygments not found. Using basic Markdown rendering.")
class LineNumberArea(QWidget):
    def __init__(self, editor):
        super().__init__(editor)
        self.editor = editor
    def sizeHint(self):
        return QSize(self.editor.line_number_area_width(), 0)
    def paintEvent(self, event):
        self.editor.line_number_area_paint_event(event)
class MarkdownHighlighter(QSyntaxHighlighter):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.highlighting_rules = []
        color_fg = QColor("#D4D4D4")
        color_bg = QColor("#1E1E1E")
        color_heading = QColor("#569CD6")
        color_bold = QColor("#D19A66")
        color_italic = QColor("#B5CEA8")
        color_list = QColor("#CE9178")
        color_link = QColor("#4EC9B0")
        color_code = QColor("#DCDCAA")
        color_quote = QColor("#6A9955")
        h1_format = QTextCharFormat()
        h1_format.setFontWeight(QFont.Weight.Bold)
        h1_format.setFontPointSize(20)
        h1_format.setForeground(color_heading)
        self.highlighting_rules.append((QRegularExpression(r"^# .+$"), h1_format))
        h2_format = QTextCharFormat()
        h2_format.setFontWeight(QFont.Weight.Bold)
        h2_format.setFontPointSize(16)
        h2_format.setForeground(color_heading)
        self.highlighting_rules.append((QRegularExpression(r"^## .+$"), h2_format))
        h3_format = QTextCharFormat()
        h3_format.setFontWeight(QFont.Weight.Bold)
        h3_format.setFontPointSize(14)
        h3_format.setForeground(color_heading)
        self.highlighting_rules.append((QRegularExpression(r"^### .+$"), h3_format))
        h4_format = QTextCharFormat()
        h4_format.setFontWeight(QFont.Weight.Bold)
        h4_format.setFontPointSize(12)
        h4_format.setForeground(color_heading)
        self.highlighting_rules.append((QRegularExpression(r"^#### .+$"), h4_format))
        bold_format = QTextCharFormat()
        bold_format.setFontWeight(QFont.Weight.Bold)
        bold_format.setForeground(color_bold)
        self.highlighting_rules.append((QRegularExpression(r"\*\*.+?\*\*"), bold_format))
        self.highlighting_rules.append((QRegularExpression(r"__.+?__"), bold_format))
        italic_format = QTextCharFormat()
        italic_format.setFontItalic(True)
        italic_format.setForeground(color_italic)
        self.highlighting_rules.append((QRegularExpression(r"\*.+?\*"), italic_format))
        self.highlighting_rules.append((QRegularExpression(r"_.+?_"), italic_format))
        list_format = QTextCharFormat()
        list_format.setForeground(color_list)
        list_format.setFontWeight(QFont.Weight.Bold)
        self.highlighting_rules.append((QRegularExpression(r"^[\*\-\+] "), list_format))
        self.highlighting_rules.append((QRegularExpression(r"^\d+\. "), list_format))
        link_format = QTextCharFormat()
        link_format.setForeground(color_link)
        link_format.setUnderlineStyle(QTextCharFormat.UnderlineStyle.SingleUnderline)
        self.highlighting_rules.append((QRegularExpression(r"\[.+?\]\(.+?\)"), link_format))
        code_format = QTextCharFormat()
        code_format.setFontFamily("Consolas")
        code_format.setBackground(QColor("#262626"))
        code_format.setForeground(color_code)
        self.highlighting_rules.append((QRegularExpression(r"`[^`]+`"), code_format))
        code_block_format = QTextCharFormat()
        code_block_format.setFontFamily("Consolas")
        code_block_format.setBackground(QColor("#23272E"))
        code_block_format.setForeground(color_code)
        self.code_block_pattern = QRegularExpression(r"```.*$")
        self.code_block_end_pattern = QRegularExpression(r"```$")
        self.code_block_format = code_block_format
        quote_format = QTextCharFormat()
        quote_format.setForeground(color_quote)
        quote_format.setFontItalic(True)
        self.highlighting_rules.append((QRegularExpression(r"^> .+$"), quote_format))
    def highlightBlock(self, text):
        previous_block_state = self.previousBlockState()
        if previous_block_state == -1:
            previous_block_state = 0
        if previous_block_state == 1:
            self.setFormat(0, len(text), self.code_block_format)
            match = self.code_block_end_pattern.match(text)
            if match.hasMatch():
                self.setCurrentBlockState(0)
            else:
                self.setCurrentBlockState(1)
            return
        match = self.code_block_pattern.match(text)
        if match.hasMatch():
            self.setFormat(0, len(text), self.code_block_format)
            self.setCurrentBlockState(1)
            return
        for pattern, format in self.highlighting_rules:
            match_iterator = pattern.globalMatch(text)
            while match_iterator.hasNext():
                match = match_iterator.next()
                self.setFormat(match.capturedStart(), match.capturedLength(), format)
class MarkdownRenderer:
    def __init__(self):
        self.md = None
        if MARKDOWN_IT_AVAILABLE:
            self.md = MarkdownIt("commonmark", {"html": True, "linkify": True, "typographer": True})
            def highlight_code(code, lang, attrs):
                if not lang:
                    return f'<pre><code>{code}</code></pre>'
                try:
                    lexer = get_lexer_by_name(lang, stripall=True)
                    formatter = HtmlFormatter(style="default", noclasses=True)
                    return f'<div class="code-block">{highlight(code, lexer, formatter)}</div>'
                except:
                    return f'<pre><code>{code}</code></pre>'
            self.md.options.highlight = highlight_code
    def render(self, text):
        if MARKDOWN_IT_AVAILABLE and self.md:
            html = self.md.render(text)
            return self._wrap_html(html)
        else:
            html = self._basic_render(text)
            return self._wrap_html(html)
    def _basic_render(self, text):
        text = re.sub(r'^# (.+)$', r'<h1>\1</h1>', text, flags=re.MULTILINE)
        text = re.sub(r'^## (.+)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
        text = re.sub(r'^### (.+)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
        text = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', text, flags=re.MULTILINE)
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'__(.+?)__', r'<strong>\1</strong>', text)
        text = re.sub(r'\*(.+?)\*', r'<em>\1</em>', text)
        text = re.sub(r'_(.+?)_', r'<em>\1</em>', text)
        text = re.sub(r'^\* (.+)$', r'<ul><li>\1</li></ul>', text, flags=re.MULTILINE)
        text = re.sub(r'^(- .+)$', r'<ul><li>\1</li></ul>', text, flags=re.MULTILINE)
        text = re.sub(r'^(\d+)\. (.+)$', r'<ol><li>\2</li></ol>', text, flags=re.MULTILINE)
        text = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', text)
        text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
        text = re.sub(r'```(.+?)```', r'<pre><code>\1</code></pre>', text, flags=re.DOTALL)
        text = re.sub(r'^> (.+)$', r'<blockquote>\1</blockquote>', text, flags=re.MULTILINE)
        paragraphs = []
        for line in text.split('\n'):
            if line.strip() and not line.startswith('<'):
                paragraphs.append(f'<p>{line}</p>')
            else:
                paragraphs.append(line)
        return '\n'.join(paragraphs)
    def _wrap_html(self, html):
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
                    line-height: 1.6;
                    padding: 20px;
                    max-width: 800px;
                    margin: 0 auto;
                }}
                h1, h2, h3, h4, h5, h6 {{
                    margin-top: 24px;
                    margin-bottom: 16px;
                    font-weight: 600;
                    color: #0366d6;
                }}
                h1 {{ font-size: 2em; padding-bottom: .3em; border-bottom: 1px solid #eaecef; }}
                h2 {{ font-size: 1.5em; padding-bottom: .3em; border-bottom: 1px solid #eaecef; }}
                h3 {{ font-size: 1.25em; }}
                h4 {{ font-size: 1em; }}
                p, blockquote, ul, ol, table {{
                    margin-bottom: 16px;
                }}
                code {{
                    font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, Courier, monospace;
                    padding: 0.2em 0.4em;
                    margin: 0;
                    font-size: 85%;
                    background-color: rgba(27, 31, 35, 0.05);
                    border-radius: 3px;
                }}
                pre {{
                    font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, Courier, monospace;
                    padding: 16px;
                    overflow: auto;
                    font-size: 85%;
                    line-height: 1.45;
                    background-color: #f6f8fa;
                    border-radius: 3px;
                }}
                pre code {{
                    background-color: transparent;
                    padding: 0;
                    margin: 0;
                    font-size: 100%;
                    word-break: normal;
                    white-space: pre;
                    border: 0;
                }}
                blockquote {{
                    padding: 0 1em;
                    color: #6a737d;
                    border-left: 0.25em solid #dfe2e5;
                }}
                ul, ol {{
                    padding-left: 2em;
                }}
                a {{
                    color: #0366d6;
                    text-decoration: none;
                }}
                a:hover {{
                    text-decoration: underline;
                }}
                table {{
                    border-spacing: 0;
                    border-collapse: collapse;
                    width: 100%;
                    overflow: auto;
                }}
                table th, table td {{
                    padding: 6px 13px;
                    border: 1px solid #dfe2e5;
                }}
                table tr {{
                    background-color: #fff;
                    border-top: 1px solid #c6cbd1;
                }}
                table tr:nth-child(2n) {{
                    background-color: #f6f8fa;
                }}
                img {{
                    max-width: 100%;
                }}
                .code-block {{
                    margin-bottom: 16px;
                }}
            </style>
        </head>
        <body>
            {html}
        </body>
        </html>
        """
class FindReplaceDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.last_match = None
        self.setWindowIcon(QIcon("Markdown_Editor.ico"))
        self.init_ui()
    def init_ui(self):
        self.setWindowTitle("Поиск и замена")
        self.setMinimumWidth(400)
        layout = QGridLayout()
        layout.addWidget(QLabel("Найти:"), 0, 0)
        self.find_input = QLineEdit()
        layout.addWidget(self.find_input, 0, 1)
        layout.addWidget(QLabel("Заменить на:"), 1, 0)
        self.replace_input = QLineEdit()
        layout.addWidget(self.replace_input, 1, 1)
        options_group = QGroupBox("Опции")
        options_layout = QVBoxLayout()
        self.case_sensitive = QCheckBox("Учитывать регистр")
        options_layout.addWidget(self.case_sensitive)
        self.whole_words = QCheckBox("Только целые слова")
        options_layout.addWidget(self.whole_words)
        options_group.setLayout(options_layout)
        layout.addWidget(options_group, 2, 0, 1, 2)
        button_layout = QGridLayout()
        find_button = QPushButton("Найти")
        find_button.clicked.connect(self.find)
        button_layout.addWidget(find_button, 0, 0)
        find_next_button = QPushButton("Найти далее")
        find_next_button.clicked.connect(self.find_next)
        button_layout.addWidget(find_next_button, 0, 1)
        replace_button = QPushButton("Заменить")
        replace_button.clicked.connect(self.replace)
        button_layout.addWidget(replace_button, 1, 0)
        replace_all_button = QPushButton("Заменить все")
        replace_all_button.clicked.connect(self.replace_all)
        button_layout.addWidget(replace_all_button, 1, 1)
        close_button = QPushButton("Закрыть")
        close_button.clicked.connect(self.close)
        button_layout.addWidget(close_button, 2, 0, 1, 2)
        layout.addLayout(button_layout, 3, 0, 1, 2)
        self.setLayout(layout)
    def find(self):
        editor = self.parent.editor
        text = self.find_input.text()
        if not text:
            return
        flags = QTextDocument.FindFlag(0)
        if self.case_sensitive.isChecked():
            flags |= QTextDocument.FindFlag.FindCaseSensitively
        if self.whole_words.isChecked():
            flags |= QTextDocument.FindFlag.FindWholeWords
        cursor = editor.textCursor()
        if cursor.hasSelection():
            cursor.setPosition(cursor.selectionStart())
            editor.setTextCursor(cursor)
        found = editor.find(text, flags)
        if not found:
            cursor.movePosition(QTextCursor.MoveOperation.Start)
            editor.setTextCursor(cursor)
            found = editor.find(text, flags)
        if found:
            self.last_match = True
        else:
            self.last_match = False
            QMessageBox.information(self, "Поиск", f"Текст '{text}' не найден")
    def find_next(self):
        if self.last_match:
            self.find()
        else:
            self.find()
    def replace(self):
        editor = self.parent.editor
        if editor.textCursor().hasSelection():
            editor.textCursor().insertText(self.replace_input.text())
            self.find()
    def replace_all(self):
        editor = self.parent.editor
        text = self.find_input.text()
        replace_text = self.replace_input.text()
        if not text:
            return
        cursor = editor.textCursor()
        cursor_position = cursor.position()
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        editor.setTextCursor(cursor)
        flags = QTextDocument.FindFlag(0)
        if self.case_sensitive.isChecked():
            flags |= QTextDocument.FindFlag.FindCaseSensitively
        if self.whole_words.isChecked():
            flags |= QTextDocument.FindFlag.FindWholeWords
        count = 0
        while editor.find(text, flags):
            cursor = editor.textCursor()
            cursor.insertText(replace_text)
            count += 1
        cursor.setPosition(cursor_position)
        editor.setTextCursor(cursor)
        QMessageBox.information(self, "Замена", f"Заменено {count} вхождений")
class FileTreeView(QTreeView):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.init_ui()
    def init_ui(self):
        self.model = QFileSystemModel()
        self.model.setReadOnly(False)
        self.model.setNameFilters(["*.md", "*.txt", "*.html", "*.docx"])
        self.model.setNameFilterDisables(False)
        self.set_root_directory(os.path.expanduser("~"))
        self.setModel(self.model)
        self.setAnimated(False)
        self.setIndentation(20)
        self.setSortingEnabled(True)
        self.sortByColumn(0, Qt.SortOrder.AscendingOrder)
        for column in range(1, self.model.columnCount()):
            self.hideColumn(column)
        self.doubleClicked.connect(self.on_double_click)
        self.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.customContextMenuRequested.connect(self.show_context_menu)
    def set_root_directory(self, path):
        self.model.setRootPath(path)
        self.setRootIndex(self.model.index(path))
    def on_double_click(self, index):
        path = self.model.filePath(index)
        if QFileInfo(path).isFile():
            self.parent.load_file(path)
    def show_context_menu(self, position):
        index = self.indexAt(position)
        if not index.isValid():
            return
        path = self.model.filePath(index)
        info = QFileInfo(path)
        menu = QMenu()
        if info.isFile():
            open_action = QAction("Открыть", self)
            open_action.triggered.connect(lambda: self.parent.load_file(path))
            menu.addAction(open_action)
        if info.isDir():
            new_file_action = QAction("Новый файл", self)
            new_file_action.triggered.connect(lambda: self.create_new_file(path))
            menu.addAction(new_file_action)
        menu.addSeparator()
        rename_action = QAction("Переименовать", self)
        rename_action.triggered.connect(lambda: self.rename_item(index))
        menu.addAction(rename_action)
        delete_action = QAction("Удалить", self)
        delete_action.triggered.connect(lambda: self.delete_item(index))
        menu.addAction(delete_action)
        menu.exec_(self.viewport().mapToGlobal(position))
    def create_new_file(self, directory):
        file_name, ok = QInputDialog.getText(
            self, "Новый файл", "Имя файла:", text="новый_файл.md"
        )
        if ok and file_name:
            file_path = os.path.join(directory, file_name)
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("")
                self.parent.load_file(file_path)
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось создать файл: {str(e)}")
    def rename_item(self, index):
        path = self.model.filePath(index)
        info = QFileInfo(path)
        old_name = info.fileName()
        new_name, ok = QInputDialog.getText(
            self, "Переименовать", "Новое имя:", text=old_name
        )
        if ok and new_name and new_name != old_name:
            new_path = os.path.join(os.path.dirname(path), new_name)
            try:
                os.rename(path, new_path)
                if self.parent.current_file == path:
                    self.parent.current_file = new_path
                    self.parent.setWindowTitle(f"Markdown Editor - {new_name}")
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось переименовать: {str(e)}")
    def delete_item(self, index):
        path = self.model.filePath(index)
        info = QFileInfo(path)
        name = info.fileName()
        reply = QMessageBox.question(
            self, "Удаление",
            f"Вы уверены, что хотите удалить {name}?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            try:
                if info.isDir():
                    import shutil
                    shutil.rmtree(path)
                else:
                    os.remove(path)
                    if self.parent.current_file == path:
                        self.parent.new_file()
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось удалить: {str(e)}")
class MarkdownEditorWidget(QPlainTextEdit):
    def __init__(self, parent=None):
        super().__init__(parent)
        font = QFont("Consolas", 13)
        self.setFont(font)
        self.setTabStopDistance(48)
        self.line_number_area = LineNumberArea(self)
        self.highlighter = MarkdownHighlighter(self.document())
        self.blockCountChanged.connect(self.update_line_number_area_width)
        self.updateRequest.connect(self.update_line_number_area)
        self.cursorPositionChanged.connect(self.highlight_current_line)
        self.cursorPositionChanged.connect(self.highlight_matching_bracket)
        self.update_line_number_area_width(0)
        self.highlight_current_line()
        self.setPlaceholderText("Введите Markdown текст здесь...")
        self.setStyleSheet("background-color: #1E1E1E; color: #D4D4D4; border-radius: 10px; border: none;")
        QShortcut(QKeySequence("Ctrl+B"), self, activated=self.shortcut_bold)
        QShortcut(QKeySequence("Ctrl+I"), self, activated=self.shortcut_italic)
        QShortcut(QKeySequence("Ctrl+K"), self, activated=self.shortcut_link)
        QShortcut(QKeySequence("Ctrl+Shift+C"), self, activated=self.shortcut_code)
        QShortcut(QKeySequence("Ctrl+Shift+L"), self, activated=self.shortcut_list)
    def shortcut_bold(self):
        self._wrap_selection("**", "**")
    def shortcut_italic(self):
        self._wrap_selection("*", "*")
    def shortcut_link(self):
        self._wrap_selection("[", "](url)")
    def shortcut_code(self):
        self._wrap_selection("`", "`")
    def shortcut_list(self):
        cursor = self.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.StartOfLine)
        cursor.insertText("- ")
        self.setTextCursor(cursor)
    def _wrap_selection(self, left, right):
        cursor = self.textCursor()
        if cursor.hasSelection():
            selected = cursor.selectedText()
            cursor.insertText(f"{left}{selected}{right}")
        else:
            cursor.insertText(f"{left}{right}")
            cursor.movePosition(QTextCursor.MoveOperation.Left, QTextCursor.MoveMode.MoveAnchor, len(right))
            self.setTextCursor(cursor)
    def highlight_matching_bracket(self):
        extra = []
        pairs = {'(': ')', '[': ']', '{': '}', '"': '"', "'": "'", '`': '`', '*': '*', '_': '_'}
        cursor = self.textCursor()
        pos = cursor.position()
        doc = self.document().toPlainText()
        if pos > 0 and pos <= len(doc):
            char = doc[pos-1]
            if char in pairs:
                match = self._find_matching(doc, pos-1, char, pairs[char])
                if match is not None:
                    sel1 = QTextEdit.ExtraSelection()
                    sel1.cursor = self.textCursor()
                    sel1.cursor.setPosition(pos-1)
                    sel1.cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.KeepAnchor, 1)
                    sel1.format.setBackground(QColor("#39C5BB"))
                    sel2 = QTextEdit.ExtraSelection()
                    sel2.cursor = self.textCursor()
                    sel2.cursor.setPosition(match)
                    sel2.cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.KeepAnchor, 1)
                    sel2.format.setBackground(QColor("#39C5BB"))
                    extra.extend([sel1, sel2])
        self.setExtraSelections(self.extraSelections() + extra)
    def _find_matching(self, text, pos, left, right):
        if left == right:
            for i in range(pos+1, len(text)):
                if text[i] == right:
                    return i
            for i in range(pos-2, -1, -1):
                if text[i] == left:
                    return i
            return None
        if left in '([{':
            depth = 1
            for i in range(pos+1, len(text)):
                if text[i] == left:
                    depth += 1
                elif text[i] == right:
                    depth -= 1
                    if depth == 0:
                        return i
        if right in ')]}':
            depth = 1
            for i in range(pos-2, -1, -1):
                if text[i] == right:
                    depth += 1
                elif text[i] == left:
                    depth -= 1
                    if depth == 0:
                        return i
        return None
    def line_number_area_width(self):
        digits = 1
        max_num = max(1, self.blockCount())
        while max_num >= 10:
            max_num /= 10
            digits += 1
        space = 3 + self.fontMetrics().horizontalAdvance('9') * digits
        return space
    def update_line_number_area_width(self, _):
        self.setViewportMargins(self.line_number_area_width(), 0, 0, 0)
    def update_line_number_area(self, rect, dy):
        if dy:
            self.line_number_area.scroll(0, dy)
        else:
            self.line_number_area.update(0, rect.y(), self.line_number_area.width(), rect.height())
        if rect.contains(self.viewport().rect()):
            self.update_line_number_area_width(0)
    def resizeEvent(self, event):
        super().resizeEvent(event)
        cr = self.contentsRect()
        self.line_number_area.setGeometry(QRect(cr.left(), cr.top(), self.line_number_area_width(), cr.height()))
    def line_number_area_paint_event(self, event):
        painter = QPainter(self.line_number_area)
        painter.fillRect(event.rect(), QColor("#23272E"))
        block = self.firstVisibleBlock()
        block_number = block.blockNumber()
        top = self.blockBoundingGeometry(block).translated(self.contentOffset()).top()
        bottom = top + self.blockBoundingRect(block).height()
        while block.isValid() and top <= event.rect().bottom():
            if block.isVisible() and bottom >= event.rect().top():
                number = str(block_number + 1)
                painter.setPen(QColor("#444B53"))
                rect = QRect(0, int(top), self.line_number_area.width(), self.fontMetrics().height())
                painter.drawText(rect, Qt.AlignmentFlag.AlignRight, number)
            block = block.next()
            top = bottom
            bottom = top + self.blockBoundingRect(block).height()
            block_number += 1
    def highlight_current_line(self):
        extra_selections = []
        if not self.isReadOnly():
            selection = QTextEdit.ExtraSelection()
            line_color = QColor("#23272E").lighter(120)
            selection.format.setBackground(line_color)
            selection.format.setProperty(QTextFormat.Property.FullWidthSelection, True)
            selection.cursor = self.textCursor()
            selection.cursor.clearSelection()
            extra_selections.append(selection)
        self.setExtraSelections(extra_selections)
class MarkdownEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Markdown Editor")
        self.setWindowIcon(QIcon("Markdown_Editor.ico"))
        self.setMinimumSize(800, 600)
        self.current_file = None
        self.file_changed = False
        self.settings = QSettings("MarkdownEditor", "MarkdownEditor")
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.main_layout = QHBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.splitter = QSplitter(Qt.Orientation.Horizontal)
        self.editor = MarkdownEditorWidget()
        self.preview = QTextBrowser()
        self.preview.setOpenExternalLinks(True)
        self.markdown_renderer = MarkdownRenderer()
        self.splitter.addWidget(self.editor)
        self.splitter.addWidget(self.preview)
        self.splitter.setSizes([400, 400])  
        self.main_layout.addWidget(self.splitter)
        self.create_file_tree()
        self.menu_bar = QMenuBar(self)
        self.setMenuBar(self.menu_bar)
        self.create_menu()
        self.create_toolbar()
        self.create_statusbar()
        self.editor.textChanged.connect(self.update_preview)
        self.editor.textChanged.connect(self.handle_text_changed)
        self.autosave_timer = QTimer(self)
        self.autosave_timer.timeout.connect(self.autosave)
        self.autosave_timer.start(10000)       
        self.load_settings()
        self.editor.setPlainText("""# Добро пожаловать в Markdown Editor!
Это **простой** редактор Markdown с *превью* в реальном времени.
## Возможности:
- Разделенный экран
- Подсветка синтаксиса
- Экспорт в разные форматы
```python
def hello_world():
    print("Hello, Markdown!")
```
[Узнать больше о Markdown](https://www.markdownguide.org/)
""")
        self.update_preview()
    def create_file_tree(self):
        self.file_tree = FileTreeView(self)
        self.file_tree_dock = QDockWidget("Файлы", self)
        self.file_tree_dock.setWidget(self.file_tree)
        self.file_tree_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.file_tree_dock)
    def create_menu(self):
        file_menu = self.menu_bar.addMenu("&Файл")
        new_action = QAction("&Новый", self)
        new_action.setShortcut("Ctrl+N")
        new_action.triggered.connect(self.new_file)
        file_menu.addAction(new_action)
        open_action = QAction("&Открыть...", self)
        open_action.setShortcut("Ctrl+O")
        open_action.triggered.connect(self.open_file)
        file_menu.addAction(open_action)
        self.recent_files_menu = QMenu("Недавние файлы", self)
        file_menu.addMenu(self.recent_files_menu)
        self.update_recent_files_menu()
        file_menu.addSeparator()
        save_action = QAction("&Сохранить", self)
        save_action.setShortcut("Ctrl+S")
        save_action.triggered.connect(self.save_file)
        file_menu.addAction(save_action)
        save_as_action = QAction("Сохранить &как...", self)
        save_as_action.setShortcut("Ctrl+Shift+S")
        save_as_action.triggered.connect(self.save_file_as)
        file_menu.addAction(save_as_action)
        file_menu.addSeparator()
        export_menu = QMenu("&Экспорт", self)
        export_html_action = QAction("Экспорт в &HTML...", self)
        export_html_action.triggered.connect(self.export_html)
        export_pdf_action = QAction("Экспорт в &PDF...", self)
        export_pdf_action.triggered.connect(self.export_pdf)
        export_docx_action = QAction("Экспорт в &DOCX...", self)
        export_docx_action.triggered.connect(self.export_docx)
        export_menu.addAction(export_html_action)
        export_menu.addAction(export_pdf_action)
        export_menu.addAction(export_docx_action)
        file_menu.addMenu(export_menu)
        file_menu.addSeparator()
        exit_action = QAction("&Выход", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)
        edit_menu = self.menu_bar.addMenu("&Правка")
        undo_action = QAction("&Отменить", self)
        undo_action.setShortcut("Ctrl+Z")
        undo_action.triggered.connect(self.editor.undo)
        edit_menu.addAction(undo_action)
        redo_action = QAction("&Повторить", self)
        redo_action.setShortcut("Ctrl+Shift+Z")
        redo_action.triggered.connect(self.editor.redo)
        edit_menu.addAction(redo_action)
        edit_menu.addSeparator()
        cut_action = QAction("&Вырезать", self)
        cut_action.setShortcut("Ctrl+X")
        cut_action.triggered.connect(self.editor.cut)
        edit_menu.addAction(cut_action)
        copy_action = QAction("&Копировать", self)
        copy_action.setShortcut("Ctrl+C")
        copy_action.triggered.connect(self.editor.copy)
        edit_menu.addAction(copy_action)
        paste_action = QAction("&Вставить", self)
        paste_action.setShortcut("Ctrl+V")
        paste_action.triggered.connect(self.editor.paste)
        edit_menu.addAction(paste_action)
        edit_menu.addSeparator()
        find_action = QAction("&Найти...", self)
        find_action.setShortcut("Ctrl+F")
        find_action.triggered.connect(self.show_find_dialog)
        edit_menu.addAction(find_action)
        replace_action = QAction("&Заменить...", self)
        replace_action.setShortcut("Ctrl+H")
        replace_action.triggered.connect(self.show_replace_dialog)
        edit_menu.addAction(replace_action)
        view_menu = self.menu_bar.addMenu("&Вид")
        editor_only_action = QAction("&Только редактор", self)
        editor_only_action.triggered.connect(self.show_editor_only)
        view_menu.addAction(editor_only_action)
        preview_only_action = QAction("&Только превью", self)
        preview_only_action.triggered.connect(self.show_preview_only)
        view_menu.addAction(preview_only_action)
        split_view_action = QAction("&Разделенный вид", self)
        split_view_action.triggered.connect(self.show_split_view)
        view_menu.addAction(split_view_action)
        view_menu.addSeparator()
        file_tree_action = QAction("Дерево файлов", self)
        file_tree_action.setCheckable(True)
        file_tree_action.setChecked(True)
        file_tree_action.triggered.connect(self.toggle_file_tree)
        view_menu.addAction(file_tree_action)
        project_menu = self.menu_bar.addMenu("&Проект")
        open_dir_action = QAction("Открыть директорию...", self)
        open_dir_action.triggered.connect(self.open_directory)
        project_menu.addAction(open_dir_action)
    def create_toolbar(self):
        self.toolbar = QToolBar("Панель инструментов")
        self.toolbar.setIconSize(QSize(18, 18))
        self.toolbar.setStyleSheet("QToolBar { spacing: 2px; padding: 2px 4px; }")
        self.addToolBar(self.toolbar)
        bold_action = QAction("Ж", self)
        bold_action.setShortcut("Ctrl+B")
        bold_action.triggered.connect(lambda: self.insert_markdown_tag("**", "**"))
        self.toolbar.addAction(bold_action)
        italic_action = QAction("К", self)
        italic_action.setShortcut("Ctrl+I")
        italic_action.triggered.connect(lambda: self.insert_markdown_tag("*", "*"))
        self.toolbar.addAction(italic_action)
        h1_action = QAction("H1", self)
        h1_action.triggered.connect(lambda: self.insert_heading(1))
        self.toolbar.addAction(h1_action)
        h2_action = QAction("H2", self)
        h2_action.triggered.connect(lambda: self.insert_heading(2))
        self.toolbar.addAction(h2_action)
        h3_action = QAction("H3", self)
        h3_action.triggered.connect(lambda: self.insert_heading(3))
        self.toolbar.addAction(h3_action)
        self.toolbar.addSeparator()
        list_action = QAction("•", self)
        list_action.triggered.connect(lambda: self.insert_list())
        self.toolbar.addAction(list_action)
        numbered_list_action = QAction("1.", self)
        numbered_list_action.triggered.connect(lambda: self.insert_numbered_list())
        self.toolbar.addAction(numbered_list_action)
        self.toolbar.addSeparator()
        link_action = QAction("🔗", self)
        link_action.triggered.connect(lambda: self.insert_markdown_tag("[", "](url)"))
        self.toolbar.addAction(link_action)
        image_action = QAction("🖼", self)
        image_action.triggered.connect(lambda: self.insert_markdown_tag("![", "](url)"))
        self.toolbar.addAction(image_action)
        code_action = QAction("<>", self)
        code_action.triggered.connect(lambda: self.insert_markdown_tag("`", "`"))
        self.toolbar.addAction(code_action)
        code_block_action = QAction("БК", self)
        code_block_action.triggered.connect(lambda: self.insert_code_block())
        self.toolbar.addAction(code_block_action)
        self.toolbar.addSeparator()
        quote_action = QAction("❝", self)
        quote_action.triggered.connect(lambda: self.insert_quote())
        self.toolbar.addAction(quote_action)
        hr_action = QAction("—", self)
        hr_action.triggered.connect(lambda: self.insert_horizontal_rule())
        self.toolbar.addAction(hr_action)
    def create_statusbar(self):
        self.statusBar().showMessage("Готово")
    def update_preview(self):
        markdown_text = self.editor.toPlainText()
        html = self.markdown_renderer.render(markdown_text)
        self.preview.setHtml(html)
    def show_editor_only(self):
        self.splitter.setSizes([1, 0])
    def show_preview_only(self):
        self.splitter.setSizes([0, 1])
    def show_split_view(self):
        self.splitter.setSizes([1, 1])
    def show_about(self):
        msg = QMessageBox(self)
        msg.setWindowIcon(QIcon("Markdown_Editor.ico"))
        msg.setIcon(QMessageBox.Icon.Information)
        msg.setWindowTitle("О программе Markdown Editor")
        msg.setText("Markdown Editor - простой редактор Markdown с превью в реальном времени.\n\nВерсия: 0.1\nАвтор: Your Name\nЛицензия: MIT")
        msg.exec()
    def insert_markdown_tag(self, start_tag, end_tag):
        cursor = self.editor.textCursor()
        selected_text = cursor.selectedText()
        if selected_text:
            cursor.insertText(f"{start_tag}{selected_text}{end_tag}")
        else:
            cursor.insertText(f"{start_tag}{end_tag}")
            cursor.movePosition(QTextCursor.MoveOperation.Left, QTextCursor.MoveMode.MoveAnchor, len(end_tag))
            self.editor.setTextCursor(cursor)
        self.editor.setFocus()
    def insert_heading(self, level):
        cursor = self.editor.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.StartOfLine)
        line_text = cursor.block().text()
        if line_text.startswith('#'):
            for i in range(len(line_text)):
                if line_text[i] == '#':
                    cursor.deleteChar()
                elif line_text[i] == ' ':
                    cursor.deleteChar()
                    break
                else:
                    break
        cursor.insertText('#' * level + ' ')
        self.editor.setFocus()
    def insert_list(self):
        cursor = self.editor.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.StartOfLine)
        cursor.insertText("- ")
        self.editor.setFocus()
    def insert_numbered_list(self):
        cursor = self.editor.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.StartOfLine)
        cursor.insertText("1. ")
        self.editor.setFocus()
    def insert_code_block(self):
        cursor = self.editor.textCursor()
        cursor.insertText("```\n")
        cursor.insertText("\n```")
        cursor.movePosition(QTextCursor.MoveOperation.Up)
        cursor.movePosition(QTextCursor.MoveOperation.EndOfLine)
        self.editor.setTextCursor(cursor)
        self.editor.setFocus()
    def insert_quote(self):
        cursor = self.editor.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.StartOfLine)
        cursor.insertText("> ")
        self.editor.setFocus()
    def insert_horizontal_rule(self):
        cursor = self.editor.textCursor()
        cursor.insertText("\n---\n")
        self.editor.setFocus()
    def handle_text_changed(self):
        self.file_changed = True
        if self.current_file:
            self.setWindowTitle(f"Markdown Editor - {os.path.basename(self.current_file)} *")
        else:
            self.setWindowTitle("Markdown Editor *")
    def new_file(self):
        if self.maybe_save():
            self.editor.clear()
            self.current_file = None
            self.file_changed = False
            self.setWindowTitle("Markdown Editor")
    def open_file(self):
        if self.maybe_save():
            file_path, _ = QFileDialog.getOpenFileName(
                self, "Открыть файл", "", 
                "Markdown Files (*.md);;Text Files (*.txt);;All Files (*)"
            )
            if file_path:
                self.load_file(file_path)
    def load_file(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                self.editor.setPlainText(file.read())
            self.current_file = file_path
            self.file_changed = False
            self.setWindowTitle(f"Markdown Editor - {os.path.basename(file_path)}")
            self.statusBar().showMessage(f"Файл {os.path.basename(file_path)} загружен")
            recent_files = self.settings.value("recentFiles", [])
            if file_path in recent_files:
                recent_files.remove(file_path)
            recent_files.insert(0, file_path)
            recent_files = recent_files[:10]  
            self.settings.setValue("recentFiles", recent_files)
            self.update_recent_files_menu()
            return True
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось открыть файл: {str(e)}")
            return False
    def save_file(self):
        if not self.current_file:
            return self.save_file_as()
        return self.save_to_file(self.current_file)
    def save_file_as(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Сохранить файл", "", 
            "Markdown Files (*.md);;Text Files (*.txt);;All Files (*)"
        )
        if file_path:
            return self.save_to_file(file_path)
        return False
    def save_to_file(self, file_path):
        try:
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(self.editor.toPlainText())
            self.current_file = file_path
            self.file_changed = False
            self.setWindowTitle(f"Markdown Editor - {os.path.basename(file_path)}")
            self.statusBar().showMessage(f"Файл {os.path.basename(file_path)} сохранен")
            recent_files = self.settings.value("recentFiles", [])
            if file_path in recent_files:
                recent_files.remove(file_path)
            recent_files.insert(0, file_path)
            recent_files = recent_files[:10]  
            self.settings.setValue("recentFiles", recent_files)
            self.update_recent_files_menu()
            return True
        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось сохранить файл: {str(e)}")
            return False
    def autosave(self):
        if self.file_changed and self.current_file:
            self.save_to_file(self.current_file)
            self.statusBar().showMessage("Автосохранение выполнено", 2000)
    def export_html(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Экспорт в HTML", "", 
            "HTML Files (*.html);;All Files (*)"
        )
        if file_path:
            try:
                html = self.markdown_renderer.render(self.editor.toPlainText())
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(html)
                self.statusBar().showMessage(f"Экспорт в HTML выполнен: {os.path.basename(file_path)}")
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось экспортировать в HTML: {str(e)}")
    def export_pdf(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Экспорт в PDF", "", 
            "PDF Files (*.pdf);;All Files (*)"
        )
        if file_path:
            try:
                from PyQt6.QtPrintSupport import QPrinter
                printer = QPrinter(QPrinter.PrinterMode.HighResolution)
                printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                printer.setOutputFileName(file_path)
                self.preview.document().print_(printer)
                self.statusBar().showMessage(f"Экспорт в PDF выполнен: {os.path.basename(file_path)}")
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось экспортировать в PDF: {str(e)}")
    def export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Экспорт в DOCX", "", 
            "Word Files (*.docx);;All Files (*)"
        )
        if file_path:
            try:
                try:
                    import docx
                except ImportError:
                    QMessageBox.warning(
                        self, "Ошибка", 
                        "Для экспорта в DOCX требуется библиотека python-docx.\n"
                        "Установите ее командой: pip install python-docx"
                    )
                    return
                from docx import Document
                document = Document()
                html_content = self.preview.toHtml()
                document.add_heading("Документ Markdown", 0)
                markdown_text = self.editor.toPlainText()
                lines = markdown_text.split('\n')
                for line in lines:
                    if line.startswith('# '):
                        document.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        document.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        document.add_heading(line[4:], level=3)
                    elif line.startswith('#### '):
                        document.add_heading(line[5:], level=4)
                    elif line.startswith('- ') or line.startswith('* '):
                        document.add_paragraph(line[2:], style='List Bullet')
                    elif re.match(r'^\d+\.\s', line):
                        document.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
                    elif line.strip():
                        line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                        line = re.sub(r'\*(.*?)\*', r'\1', line)
                        line = re.sub(r'`(.*?)`', r'\1', line)
                        line = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', line)
                        document.add_paragraph(line)
                    else:
                        document.add_paragraph('')
                document.save(file_path)
                self.statusBar().showMessage(f"Экспорт в DOCX выполнен: {os.path.basename(file_path)}")
            except Exception as e:
                QMessageBox.warning(self, "Ошибка", f"Не удалось экспортировать в DOCX: {str(e)}")
    def maybe_save(self):
        if not self.file_changed:
            return True
        reply = QMessageBox.question(
            self, "Сохранение изменений",
            "Документ был изменен. Сохранить изменения?",
            QMessageBox.StandardButton.Save | QMessageBox.StandardButton.Discard | QMessageBox.StandardButton.Cancel
        )
        if reply == QMessageBox.StandardButton.Save:
            return self.save_file()
        elif reply == QMessageBox.StandardButton.Cancel:
            return False
        return True
    def closeEvent(self, event):
        if self.maybe_save():
            self.save_settings()
            event.accept()
        else:
            event.ignore()
    def load_settings(self):
        geometry = self.settings.value("geometry")
        if geometry:
            self.restoreGeometry(geometry)
        splitter_sizes = self.settings.value("splitterSizes")
        if splitter_sizes and isinstance(splitter_sizes, list) and all(isinstance(size, int) for size in splitter_sizes):
            self.splitter.setSizes(splitter_sizes)
        last_directory = self.settings.value("lastDirectory")
        if last_directory and isinstance(last_directory, str) and os.path.exists(last_directory):
            self.file_tree.set_root_directory(last_directory)
    def save_settings(self):
        self.settings.setValue("geometry", self.saveGeometry())
        self.settings.setValue("splitterSizes", self.splitter.sizes())
        self.update_recent_files_menu()
    def show_find_dialog(self):
        dialog = FindReplaceDialog(self)
        dialog.setWindowIcon(QIcon("Markdown_Editor.ico"))
        dialog.replace_input.hide()
        dialog.layout().itemAtPosition(1, 0).widget().hide()
        dialog.exec()
    def show_replace_dialog(self):
        dialog = FindReplaceDialog(self)
        dialog.setWindowIcon(QIcon("Markdown_Editor.ico"))
        dialog.exec()
    def update_recent_files_menu(self):
        self.recent_files_menu.clear()
        recent_files = self.settings.value("recentFiles", [])
        for file_path in recent_files:
            if os.path.exists(file_path):
                action = QAction(os.path.basename(file_path), self)
                action.setData(file_path)
                action.triggered.connect(self.open_recent_file)
                self.recent_files_menu.addAction(action)
    def open_recent_file(self):
        action = self.sender()
        if action:
            file_path = action.data()
            if os.path.exists(file_path):
                if self.maybe_save():
                    self.load_file(file_path)
    def toggle_file_tree(self):
        if self.file_tree_dock.isVisible():
            self.file_tree_dock.hide()
        else:
            self.file_tree_dock.show()
    def open_directory(self):
        directory = QFileDialog.getExistingDirectory(
            self, "Открыть директорию", "", QFileDialog.Option.ShowDirsOnly
        )
        if directory:
            self.file_tree.set_root_directory(directory)
            self.settings.setValue("lastDirectory", directory)
def apply_modern_dark_theme(app):
    app.setStyle("Fusion")
    font = QFont("Segoe UI", 12)
    app.setFont(font)
    qss = """
    QWidget {
        background-color: #181A1B;
        color: #E4E6EB;
        font-size: 16px;
    }
    QMainWindow, QDialog, QMenuBar, QMenu, QStatusBar {
        background-color: #181A1B;
        color: #E4E6EB;
        border: none;
    }
    QMenuBar {
        font-size: 18px;
        padding: 0 8px;
        background: #181A1B;
    }
    QMenuBar::item {
        background: transparent;
        padding: 6px 18px;
        margin: 0 2px;
        border-radius: 8px;
        min-width: 32px;
        min-height: 24px;
    }
    QMenuBar::item:selected {
        background: #23272E;
        color: #7FDBFF;
    }
    QMenu {
        font-size: 16px;
        padding: 6px 12px;
        border-radius: 8px;
    }
    QMenu::item {
        padding: 6px 18px;
        border-radius: 6px;
    }
    QMenu::item:selected {
        background-color: #31343B;
        color: #7FDBFF;
    }
    QToolBar {
        background: #23272E;
        border: none;
        spacing: 4px;
        padding: 6px 8px;
    }
    QToolButton {
        background-color: #23272E;
        color: #E4E6EB;
        border-radius: 7px;
        padding: 4px 10px;
        font-size: 14px;
        min-width: 28px;
        min-height: 28px;
        max-width: 32px;
        max-height: 32px;
        border: 1px solid #23272E;
    }
    QToolButton:hover {
        background-color: #31343B;
        color: #7FDBFF;
    }
    QToolButton:pressed {
        background-color: #1A1D22;
        color: #39C5BB;
    }
    QPushButton {
        background-color: #23272E;
        color: #E4E6EB;
        border-radius: 12px;
        padding: 10px 20px;
        font-size: 16px;
        border: 1px solid #23272E;
    }
    QPushButton:hover {
        background-color: #31343B;
        color: #7FDBFF;
    }
    QPushButton:pressed {
        background-color: #1A1D22;
        color: #39C5BB;
    }
    QLineEdit, QTextEdit, QPlainTextEdit, QTextBrowser {
        background-color: #23272E;
        color: #E4E6EB;
        border-radius: 10px;
        border: 1px solid #23272E;
        padding: 8px;
        font-size: 16px;
        selection-background-color: #31343B;
        selection-color: #7FDBFF;
    }
    QStatusBar {
        background: #181A1B;
        color: #7FDBFF;
        border: none;
    }
    QGroupBox {
        border: 1px solid #23272E;
        border-radius: 10px;
        margin-top: 10px;
    }
    QGroupBox:title {
        subcontrol-origin: margin;
        left: 10px;
        padding: 0 3px 0 3px;
    }
    QTreeView, QDockWidget {
        background-color: #181A1B;
        color: #E4E6EB;
        border-radius: 10px;
    }
    QScrollBar:vertical, QScrollBar:horizontal {
        background: #23272E;
        border-radius: 8px;
        width: 14px;
        margin: 2px;
    }
    QScrollBar::handle:vertical, QScrollBar::handle:horizontal {
        background: #31343B;
        border-radius: 8px;
        min-height: 20px;
        min-width: 20px;
    }
    QScrollBar::handle:hover {
        background: #7FDBFF;
    }
    QScrollBar::add-line, QScrollBar::sub-line {
        background: none;
        border: none;
    }
    QLabel {
        color: #E4E6EB;
        font-size: 16px;
    }
    """
    app.setStyleSheet(qss)
def main():
    app = QApplication(sys.argv)
    apply_modern_dark_theme(app)
    window = MarkdownEditor()
    window.show()
    sys.exit(app.exec())
if __name__ == "__main__":
    main() 
